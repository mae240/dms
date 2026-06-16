"""Tests fuer die Verarbeitungs-Pipeline (reine Logik + Textextraktion)."""

from __future__ import annotations

import hashlib
import io

from dms_core.enums import ProcessingStatus
from worker.extract import extract_text
from worker.tasks.processing import evaluate_blob

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_extract_text_plain() -> None:
    assert extract_text("text/plain", b"Hallo Welt") == "Hallo Welt"


def test_extract_text_unknown_returns_none() -> None:
    assert extract_text("image/png", b"\x89PNG\r\n") is None


def test_extract_text_docx() -> None:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Vertragsklausel Alpha")
    doc.add_paragraph("Vertragsklausel Beta")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text(DOCX_MIME, buf.getvalue())
    assert text is not None
    assert "Vertragsklausel Alpha" in text
    assert "Vertragsklausel Beta" in text


def test_evaluate_blob_ready_and_extracts_text() -> None:
    data = b"Dies ist ein einfacher Vertragstext."
    status, error, text = evaluate_blob(
        data, expected_hash=hashlib.sha256(data).hexdigest(), stored_mime="text/plain"
    )
    assert status == ProcessingStatus.ready
    assert error is None
    assert text == data.decode()


def test_evaluate_blob_hash_mismatch_is_failed() -> None:
    data = b"inhalt"
    status, error, _ = evaluate_blob(data, expected_hash="0" * 64, stored_mime="text/plain")
    assert status == ProcessingStatus.failed
    assert error and "SHA-256" in error


def test_evaluate_blob_mime_mismatch_is_quarantined() -> None:
    data = b"eigentlich nur text"
    status, error, _ = evaluate_blob(
        data, expected_hash=hashlib.sha256(data).hexdigest(), stored_mime="application/pdf"
    )
    assert status == ProcessingStatus.quarantined
    assert error and "MIME" in error
